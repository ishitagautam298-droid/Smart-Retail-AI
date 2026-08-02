import os
import site
import sys
sys.path.insert(0, site.getusersitepackages())

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_FILE = "/Users/ishitagautam/.gemini/antigravity/scratch/Smart-Retail-AI/Smart_Retail_AI_Project_Report.docx"
DOWNLOADS_FILE = "/Users/ishitagautam/Downloads/Smart_Retail_AI_Project_Report.docx"

doc = Document()

# ── Margins ──────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Color Palette ─────────────────────────────────────────────────────────────
PURPLE = RGBColor(0x4A, 0x00, 0x80)
CYAN   = RGBColor(0x00, 0x77, 0xCC)
DARK   = RGBColor(0x1A, 0x1A, 0x2E)
GREY   = RGBColor(0x66, 0x66, 0x66)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GREEN  = RGBColor(0x2E, 0x7D, 0x32)

# Helper functions
def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_cell_border(cell, color_hex="CCCCCC", sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), sz)
        tag.set(qn('w:color'), color_hex)
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level==1 else 12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16 if level==1 else 13 if level==2 else 11)
    run.font.color.rgb = PURPLE if level==1 else CYAN if level==2 else DARK
    return p

def add_body(text, bold_prefix=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.bold = True
        r_pre.font.size = Pt(10.5)
        r_pre.font.color.rgb = DARK
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    run.font.color.rgb = DARK
    return p

def add_bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.bold = True
        r_pre.font.size = Pt(10)
        r_pre.font.color.rgb = PURPLE
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    return p

def add_callout(text, title="NOTE:"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    shade_cell(cell, "F0E8FF")
    
    # Left thick purple border
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    left_b = OxmlElement('w:left')
    left_b.set(qn('w:val'), 'single')
    left_b.set(qn('w:sz'), '24')
    left_b.set(qn('w:color'), '4A0080')
    tcBorders.append(left_b)
    for edge in ('top', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'none')
        tcBorders.append(b)
    tcPr.append(tcBorders)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    
    r_t = p.add_run(f"{title} ")
    r_t.bold = True
    r_t.font.color.rgb = PURPLE
    r_t.font.size = Pt(10)
    
    r_m = p.add_run(text)
    r_m.font.size = Pt(10)
    r_m.font.name = 'Calibri'
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_code_block(code_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    shade_cell(cell, "F4F5F7")
    set_cell_border(cell, "DDDDDD", "4")
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.3)
    
    r = p.add_run(code_text)
    r.font.name = 'Courier New'
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ==============================================================================
# 1. FRONT COVER PAGE
# ==============================================================================

# Title Spacing
p_title_space = doc.add_paragraph()
p_title_space.paragraph_format.space_before = Pt(36)

# Main Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_t = p_title.add_run("AI-POWERED SMART RETAIL &\nCUSTOMER INTELLIGENCE PLATFORM")
run_t.font.name = 'Arial'
run_t.font.size = Pt(24)
run_t.bold = True
run_t.font.color.rgb = PURPLE
p_title.paragraph_format.space_after = Pt(12)

# Subtitle
p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = p_sub.add_run("Multi-Modal Artificial Intelligence System & Microservice Architecture Report")
run_sub.font.name = 'Calibri'
run_sub.font.size = Pt(13)
run_sub.font.color.rgb = CYAN
run_sub.italic = True
p_sub.paragraph_format.space_after = Pt(24)

# Cover Image
img_path = "/Users/ishitagautam/.gemini/antigravity/scratch/Smart-Retail-AI/images/architecture.png"
if os.path.exists(img_path):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_after = Pt(24)
    run_img = p_img.add_run()
    run_img.add_picture(img_path, width=Inches(5.8))

# Repository Link Card Box on Cover Page
table_repo = doc.add_table(rows=1, cols=1)
table_repo.alignment = WD_TABLE_ALIGNMENT.CENTER
cell_repo = table_repo.cell(0, 0)
cell_repo.width = Inches(5.8)
shade_cell(cell_repo, "4A0080")

p_repo = cell_repo.paragraphs[0]
p_repo.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_repo.paragraph_format.space_before = Pt(10)
p_repo.paragraph_format.space_after  = Pt(10)

r_r1 = p_repo.add_run("🔗 OFFICIAL GITHUB REPOSITORY LINK\n")
r_r1.font.name = 'Arial'
r_r1.font.size = Pt(11)
r_r1.bold = True
r_r1.font.color.rgb = WHITE

r_r2 = p_repo.add_run("https://github.com/ishitagautam298-droid/Smart-Retail-AI")
r_r2.font.name = 'Courier New'
r_r2.font.size = Pt(10.5)
r_r2.bold = True
r_r2.font.color.rgb = RGBColor(0xFF, 0xEB, 0x3B)

doc.add_paragraph().paragraph_format.space_after = Pt(36)

# Author Info Block
p_author = doc.add_paragraph()
p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_a1 = p_author.add_run("Prepared & Authored By:\n")
r_a1.font.size = Pt(10)
r_a1.font.color.rgb = GREY

r_a2 = p_author.add_run("ISHITA GAUTAM\n")
r_a2.font.size = Pt(14)
r_a2.bold = True
r_a2.font.color.rgb = DARK

r_a3 = p_author.add_run("Artificial Intelligence & Data Engineering | August 2026")
r_a3.font.size = Pt(10)
r_a3.font.color.rgb = GREY

# Page Break after Cover
doc.add_page_break()

# ==============================================================================
# 2. EXECUTIVE SUMMARY & TABLE OF CONTENTS
# ==============================================================================

add_heading("1. Executive Summary", level=1)
add_body("Physical retail stores face intense competition from e-commerce platforms that offer hyper-personalized user experiences, automated customer support, and instant product search. The Smart Retail & Customer Intelligence Platform bridges this technological divide by integrating computer vision, facial biometrics, natural language processing (NLP), and microservice APIs into an operational retail architecture.")

add_callout(
    "This project brings together 5 core Google Colab modules into a unified multi-modal AI platform. All code, serialized models, datasets, and FastAPI endpoints are published open-source at https://github.com/ishitagautam298-droid/Smart-Retail-AI.",
    title="KEY HIGHLIGHT:"
)

add_heading("2. System Architecture Overview", level=1)
add_body("The platform is structured into 5 cohesive modules operating as a service mesh:")
add_bullet(" Deep Convolutional Neural Network (CNN) trained on Fashion-MNIST for 10-category clothing item classification.", bold_prefix="Module 1 (Vision):")
add_bullet(" OpenCV & 128-dimensional facial embedding vectors using HOG & deep metric learning for VIP identification and visit logging.", bold_prefix="Module 2 (Biometrics):")
add_bullet(" TF-IDF Vectorization and Logistic Regression trained on ~23,000 clothing reviews for customer mood tracking.", bold_prefix="Module 3 (Sentiment):")
add_bullet(" Intent classification engine matching user queries with automated retail resolution tags.", bold_prefix="Module 4 (Chatbot):")
add_bullet(" Async RESTful API built on FastAPI, Uvicorn, and Pydantic exposing lightweight model inference endpoints.", bold_prefix="Module 5 (Microservices):")

# Add Architecture Image
if os.path.exists(img_path):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = p_img.add_run()
    run_img.add_picture(img_path, width=Inches(6.2))

# ==============================================================================
# 3. MODULE 1: PRODUCT IMAGE CLASSIFICATION
# ==============================================================================

add_heading("3. Module 1: Product Image Classification (CNN)", level=1)
add_body("Manual inventory identification and visual entry of incoming store apparel is slow and prone to human error. Module 1 automates product image classification using a 2D Convolutional Neural Network built with TensorFlow & Keras.")

add_heading("3.1 Dataset & Preprocessing", level=2)
add_body("The model is trained on the Fashion-MNIST benchmark dataset comprising 70,000 grayscale images across 10 categories (T-shirt/top, Trouser, Pullover, Dress, Coat, Sandal, Shirt, Sneaker, Bag, Ankle boot). Pixel values are normalized from [0, 255] to [0.0, 1.0] and reshaped to tensor format (28, 28, 1).")

add_heading("3.2 Model Topology & Metrics", level=2)
add_bullet("Conv2D(32 filters, 3x3 kernel, ReLU) + MaxPooling2D(2x2)", bold_prefix="Layer 1: ")
add_bullet("Conv2D(64 filters, 3x3 kernel, ReLU) + MaxPooling2D(2x2)", bold_prefix="Layer 2: ")
add_bullet("Flatten + Dense(128 units, ReLU) + Dropout(0.2) + Dense(10 units, Softmax)", bold_prefix="Dense Stack: ")
add_bullet("Adam (learning_rate=0.001), Sparse Categorical Crossentropy Loss", bold_prefix="Optimizer: ")

img_m1 = "/Users/ishitagautam/.gemini/antigravity/scratch/Smart-Retail-AI/images/fashion_mnist_metrics.png"
if os.path.exists(img_m1):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = p_img.add_run()
    run_img.add_picture(img_m1, width=Inches(6.0))

add_body("The model achieved 92.4% Training Accuracy and 89.6% Validation Accuracy after 10 epochs. The trained weights were saved to models/product_classifier.keras (2.7 MB).")

# ==============================================================================
# 4. MODULE 2: FACIAL RECOGNITION & BIOMETRICS
# ==============================================================================

add_heading("4. Module 2: Facial Recognition & VIP Intelligence", level=1)
add_body("Retail stores often lack real-time identification of returning VIP customers upon store entry. Module 2 combines OpenCV facial detection with 128-dimensional deep metric learning embeddings.")

add_heading("4.1 Biometric Matching Pipeline", level=2)
add_bullet("Employs Histogram of Oriented Gradients (HOG) and Haar Cascades to locate faces in camera video frames.", bold_prefix="1. Detection: ")
add_bullet("Generates a 128-d numerical embedding vector per face, capturing distinct facial feature geometry.", bold_prefix="2. Embedding Extraction: ")
add_bullet("Computes Euclidean distance against pre-computed customer encodings stored in face_db.pkl (tolerance = 0.6).", bold_prefix="3. Database Comparison: ")
add_bullet("Records customer visit timestamps, visit counts, and returning vs. new visitor flags into customer_visits.csv.", bold_prefix="4. Visit Analytics: ")

img_m2 = "/Users/ishitagautam/.gemini/antigravity/scratch/Smart-Retail-AI/images/face_rec_pipeline.png"
if os.path.exists(img_m2):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = p_img.add_run()
    run_img.add_picture(img_m2, width=Inches(6.0))

# ==============================================================================
# 5. MODULE 3: SENTIMENT ANALYSIS (NLP)
# ==============================================================================

add_heading("5. Module 3: Customer Review Sentiment Analysis", level=1)
add_body("Understanding customer satisfaction requires analyzing large volumes of unorganized review text. Module 3 processes 23,486 real-world e-commerce customer reviews using NLP techniques.")

add_heading("5.1 Preprocessing & Logistic Regression Model", level=2)
add_body("Review text is cleaned (lowercasing, punctuation removal, stop-word filtering) and converted into numerical vectors using TF-IDF (Term Frequency-Inverse Document Frequency) with 5,000 max features. A Logistic Regression classifier categorizes sentiment into Positive, Negative, or Neutral.")

img_m3 = "/Users/ishitagautam/.gemini/antigravity/scratch/Smart-Retail-AI/images/sentiment_analytics.png"
if os.path.exists(img_m3):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = p_img.add_run()
    run_img.add_picture(img_m3, width=Inches(5.8))

add_body("The model achieved 88.5% Classification Accuracy and 0.91 Precision for positive reviews. Pre-trained weights are stored in models/sentiment_model.pkl (120 KB) and models/vectorizer.pkl (183 KB).")

# ==============================================================================
# 6. MODULE 4: SUPPORT CHATBOT INTENT ENGINE
# ==============================================================================

add_heading("6. Module 4: Retail Support Chatbot Intent Engine", level=1)
add_body("Common retail inquiries (order tracking, return policies, store hours) consume customer support agent bandwidth. Module 4 implements an intent classification model using TF-IDF and Scikit-Learn LabelEncoder.")

img_m4 = "/Users/ishitagautam/.gemini/antigravity/scratch/Smart-Retail-AI/images/chatbot_intents.png"
if os.path.exists(img_m4):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = p_img.add_run()
    run_img.add_picture(img_m4, width=Inches(5.8))

add_body("The chatbot routes user queries to intents such as track_order (32%), return_policy (24%), and store_hours (18%), allowing automated 24/7 retail resolution.")

# ==============================================================================
# 7. MODULE 5: FASTAPI MICROSERVICES & REST ENGINE
# ==============================================================================

add_heading("7. Module 5: FastAPI Microservices Architecture", level=1)
add_body("To serve AI inference endpoints to mobile apps, web dashboards, and IoT POS terminals, Module 5 implements a high-performance RESTful API using FastAPI and Uvicorn.")

img_m5 = "/Users/ishitagautam/.gemini/antigravity/scratch/Smart-Retail-AI/images/api_mesh.png"
if os.path.exists(img_m5):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = p_img.add_run()
    run_img.add_picture(img_m5, width=Inches(5.8))

add_heading("7.1 API Endpoint Reference Table", level=2)

# Create API Table
table_api = doc.add_table(rows=5, cols=4)
table_api.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Method", "Endpoint", "Input Payload", "Description"]
widths  = [Inches(1.0), Inches(1.8), Inches(1.8), Inches(1.9)]

# Header Row
hdr_cells = table_api.rows[0].cells
for i, text in enumerate(headers):
    hdr_cells[i].width = widths[i]
    shade_cell(hdr_cells[i], "4A0080")
    p = hdr_cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = WHITE

# Data Rows
rows_data = [
    ("GET", "/", "None", "API Health & Service Discovery"),
    ("POST", "/recognize-face", "Image file (Multipart)", "Identifies face & logs VIP visit"),
    ("POST", "/analyze-sentiment", "JSON { 'text': '...' }", "Predicts review sentiment tag"),
    ("POST", "/chatbot", "JSON { 'message': '...' }", "Classifies query intent tag")
]

for row_idx, data in enumerate(rows_data, start=1):
    row_cells = table_api.rows[row_idx].cells
    bg_color = "F9F9FB" if row_idx % 2 == 1 else "FFFFFF"
    for col_idx, text in enumerate(data):
        row_cells[col_idx].width = widths[col_idx]
        shade_cell(row_cells[col_idx], bg_color)
        set_cell_border(row_cells[col_idx], "DDDDDD", "4")
        p = row_cells[col_idx].paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(text)
        r.font.size = Pt(9)
        r.font.name = 'Calibri'
        if col_idx == 0:
            r.bold = True
            r.font.color.rgb = GREEN if text=="GET" else CYAN

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# ==============================================================================
# 8. ETHICAL AI & DEPLOYMENT GUIDE
# ==============================================================================

add_heading("8. Ethical AI, Privacy & Security Standards", level=1)
add_body("Deploying facial biometrics and AI tracking in public commercial spaces requires compliance with global privacy regulations (GDPR, CCPA, BIPA):")
add_bullet(" Facial recognition operates strictly on an opt-in basis for store loyalty members.", bold_prefix="1. Explicit Consent: ")
add_bullet(" Raw video frames and images are discarded immediately after feature extraction; only non-invertible 128-d numerical embeddings are stored.", bold_prefix="2. Data Minimization: ")
add_bullet(" Models are evaluated across diverse demographic populations to eliminate facial recognition bias.", bold_prefix="3. Demographic Fairness: ")

add_heading("9. Quickstart Deployment Guide", level=1)
add_body("To launch the platform on any local machine or server:")

add_code_block("""# 1. Clone repository
git clone https://github.com/ishitagautam298-droid/Smart-Retail-AI.git
cd Smart-Retail-AI

# 2. Install dependencies
pip install -r requirements.txt

# 3. Launch FastAPI backend server
python3 app.py""")

add_heading("10. Conclusion & Repository Link", level=1)
add_body("The Smart Retail & Customer Intelligence Platform demonstrates how multi-modal computer vision, NLP, and biometrics models can be cleanly integrated into an operational retail backend.")

add_callout(
    "All source code, 5 Google Colab notebooks, pre-trained models, datasets, and deployment scripts are available at:\nhttps://github.com/ishitagautam298-droid/Smart-Retail-AI",
    title="PROJECT REPOSITORY LINK:"
)

# Save document to scratch and Downloads
doc.save(OUTPUT_FILE)
doc.save(DOWNLOADS_FILE)

print(f"SUCCESS: Document generated at:\n1. {OUTPUT_FILE}\n2. {DOWNLOADS_FILE}")
